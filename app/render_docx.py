"""Génération Word (.docx) des CV pour les deux modèles (python-docx)."""

from __future__ import annotations

import io

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.labels import labels_for
from app.render_pdf import LOGOS_DIR, MODEL_CONFIG, apply_commercial_defaults
from app.schema import CV, Language, TemplateModel

DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.lstrip("#").upper())


# Style de titres par modèle (majuscules Libens, souligné Compleo).
STYLE = {
    TemplateModel.libens: {"underline": False, "upper": True},
    TemplateModel.compleo: {"underline": True, "upper": False},
}


def _add_bottom_border(paragraph, color: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color.lstrip("#"))
    pbdr.append(bottom)
    pPr.append(pbdr)


def _run(paragraph, text, *, color=None, bold=False, size=None, italic=False):
    r = paragraph.add_run(text)
    r.font.bold = bold
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    return r


class _Builder:
    def __init__(self, cv: CV, model: TemplateModel, language: Language = Language.fr):
        self.cv = cv
        self.model = model
        self.cfg = MODEL_CONFIG[model]
        self.lab = labels_for(language)
        self.style = STYLE[model]
        self.orange = _rgb(self.cfg["orange"])
        self.blue = _rgb(self.cfg["blue"])
        # Couleur du corps : bleu pour Compleo, sombre pour Libens.
        self.body = self.blue if model == TemplateModel.compleo else DARK
        self.doc = Document()
        self._page_setup()

    def _page_setup(self):
        for s in self.doc.sections:
            s.top_margin = Cm(1.4)
            s.bottom_margin = Cm(1.4)
            s.left_margin = Cm(1.8)
            s.right_margin = Cm(1.8)
        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10)

    # -- En-tête répété (logo, + contact pour Compleo) --
    def _header(self):
        header = self.doc.sections[0].header
        logo = LOGOS_DIR / self.cfg["logo"]
        if self.model == TemplateModel.libens:
            p = header.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if logo.exists():
                p.add_run().add_picture(str(logo), height=Cm(1.5))
        else:
            table = header.add_table(rows=1, cols=2, width=Cm(17.8))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            left, right = table.rows[0].cells
            lp = left.paragraphs[0]
            if logo.exists():
                lp.add_run().add_picture(str(logo), height=Cm(1.3))
            rp = right.paragraphs[0]
            rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Contact du commercial Compleo (pas le candidat).
            if self.cv.commercial_name:
                _run(rp, self.cv.commercial_name, color=DARK, bold=True, size=9)
                rp.add_run().add_break()
            if self.cv.commercial_email:
                _run(rp, self.cv.commercial_email, color=GREY, size=8.5)
                rp.add_run().add_break()
            if self.cv.commercial_phone:
                _run(rp, self.cv.commercial_phone, color=GREY, size=8.5)
            # supprime le paragraphe vide par défaut du header
            self.doc.sections[0].header.paragraphs[0].clear()

    def _title_block(self):
        color = self.blue if self.model == TemplateModel.compleo else self.orange
        display_name = (
            self.cv.full_name
            if self.model == TemplateModel.compleo and self.cv.full_name
            else self.cv.initials
        )
        if display_name:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, display_name, color=color, bold=True, size=14)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, self.cv.title, color=color, bold=True, size=12)

    def _section(self, label: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        text = label.upper() if self.style["upper"] else label
        r = _run(p, text, color=self.orange, bold=True, size=11)
        # Soulignement collé au texte du titre (pas de trait pleine largeur).
        if self.style["underline"]:
            r.font.underline = True

    def _para(self, text, *, color=None, bold=False, space_after=4):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        _run(p, text, color=color if color is not None else self.body, bold=bold)
        return p

    def _bullets(self, items):
        for it in items:
            p = self.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            _run(p, it, color=self.body)

    def _sub(self, label: str) -> str:
        # Compleo suffixe les sous-titres d'expérience par " :".
        return f"{label} :" if self.model == TemplateModel.compleo else label

    def build(self) -> bytes:
        self._header()
        self._title_block()
        cv, lab = self.cv, self.lab

        if cv.summary:
            self._section(lab["summary"])
            for para in [x.strip() for x in cv.summary.split("\n") if x.strip()]:
                self._para(para)

        if lab["highlights"] and cv.experience_highlights:
            self._section(lab["highlights"])
            self._bullets(cv.experience_highlights)

        if cv.skill_groups:
            self._section(lab["skills"])
            for g in cv.skill_groups:
                self._para(g.category, bold=True, space_after=1)
                self._bullets(g.items)

        if cv.education:
            self._section(lab["education"])
            for e in cv.education:
                parts = [e.degree] + [x for x in (e.school, e.period) if x]
                self._para(" — ".join(parts), bold=False)

        if cv.certifications:
            self._section(lab["certs"])
            self._bullets(cv.certifications)

        if cv.languages:
            self._section(lab["langs"])
            self._bullets(cv.languages)

        if cv.experiences:
            self._section(lab["experience"])
            for x in cv.experiences:
                head = self.doc.add_paragraph()
                head.paragraph_format.space_before = Pt(6)
                comp = x.company + (f" – {x.location}" if x.location else "")
                _run(head, comp, color=self.blue, bold=True, size=10.5)
                if x.period:
                    _run(head, "    " + x.period, color=self.blue, bold=True)
                if x.role:
                    self._para(x.role, color=self.blue, bold=True, space_after=2)
                if x.context:
                    self._para(self._sub(lab["context"]), bold=True, space_after=1)
                    self._para(x.context)
                if x.responsibilities:
                    self._para(self._sub(lab["responsibilities"]), bold=True, space_after=1)
                    self._bullets(x.responsibilities)
                if x.tech_environment:
                    self._para(self._sub(lab["tech"]), bold=True, space_after=1)
                    self._para(x.tech_environment)

        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()


def render_docx(cv: CV, model: TemplateModel, language: Language = Language.fr) -> bytes:
    return _Builder(apply_commercial_defaults(cv, model), model, language).build()
