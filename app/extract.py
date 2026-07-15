"""Extraction d'un CV brut vers le schéma structuré via Claude.

Étapes :
1. Extraire le texte du fichier (PDF / DOCX / TXT).
2. Demander à Claude de remplir le schéma CV via un outil (function calling),
   ce qui garantit une sortie JSON conforme.
"""

from __future__ import annotations

import io

from anthropic import Anthropic

from app.config import settings
from app.schema import CV

_client: Anthropic | None = None


def _anthropic() -> Anthropic:
    global _client
    if _client is None:
        _client = Anthropic(api_key=settings.anthropic_api_key)
    return _client


def compute_initials(full_name: str | None) -> str | None:
    """Code anonymisé selon la règle : 1re lettre du prénom + 1re et dernière
    lettre du nom de famille. Ex : 'Moundir Jorio' -> 'MJO'.
    """
    if not full_name:
        return None
    parts = [p for p in full_name.split() if p]
    if not parts:
        return None
    first, last = parts[0], parts[-1]
    if len(parts) == 1:
        return (first[0] + first[-1]).upper()
    return (first[0] + last[0] + last[-1]).upper()


def extract_text_from_file(filename: str, data: bytes) -> str:
    """Extrait le texte brut d'un fichier CV (PDF, DOCX ou texte)."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if name.endswith(".docx"):
        from docx import Document

        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    # Fallback : texte brut
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="ignore")


_SYSTEM = (
    "Tu es un expert RH qui structure des CV pour des sociétés de conseil. "
    "À partir du texte brut d'un CV, tu remplis fidèlement le schéma fourni. "
    "Règles :\n"
    "- Ne jamais inventer d'information absente ; laisse vide si inconnu.\n"
    "- 'full_name' = nom réel complet du candidat si présent.\n"
    "- 'initials' = code anonymisé du candidat. Règle : 1re lettre du prénom "
    "+ 1re lettre du nom + dernière lettre du nom (ex: 'Moundir Jorio' -> 'MJO'). "
    "Si le nom complet est absent mais qu'un code figure déjà dans le CV "
    "(ex: 'P.B.Y'), reprends-le tel quel.\n"
    "- Conserve la langue d'origine du CV pour les contenus.\n"
    "- Regroupe les compétences par catégorie logique.\n"
    "- Pour chaque expérience, remplis company, role, period, context, "
    "responsibilities (liste) et tech_environment quand disponible.\n"
    "- Respecte l'ordre chronologique inverse (plus récent d'abord)."
)


def _cv_tool() -> dict:
    schema = CV.model_json_schema()
    return {
        "name": "emit_cv",
        "description": "Émet le CV structuré conforme au schéma.",
        "input_schema": schema,
    }


def structure_cv(raw_text: str) -> CV:
    """Appelle Claude pour transformer le texte brut en objet CV."""
    if not raw_text.strip():
        raise ValueError("Texte du CV vide : extraction impossible.")

    tool = _cv_tool()
    msg = _anthropic().messages.create(
        model=settings.claude_model,
        max_tokens=16000,
        # Adaptive thinking : Claude raisonne avant de structurer (utile pour les
        # CV denses/ambigus). tool_choice="auto" est requis pour que le thinking
        # s'active réellement ; le repli plus bas garantit qu'emit_cv a été appelé.
        thinking={"type": "adaptive"},
        system=_SYSTEM,
        tools=[tool],
        tool_choice={"type": "auto"},
        messages=[
            {
                "role": "user",
                "content": (
                    "Voici le texte brut d'un CV. Structure-le dans le schéma.\n\n"
                    "=== CV BRUT ===\n" + raw_text[:60000]
                ),
            }
        ],
    )
    for block in msg.content:
        if block.type == "tool_use" and block.name == "emit_cv":
            cv = CV.model_validate(block.input)
            # Le code anonymisé est recalculé de façon déterministe dès qu'on
            # connaît le nom complet (règle prénom[0] + nom[0] + nom[-1]).
            if cv.full_name:
                cv.initials = compute_initials(cv.full_name) or cv.initials
            return cv
    raise RuntimeError("Claude n'a pas renvoyé de CV structuré.")


def extract_cv(filename: str, data: bytes) -> CV:
    text = extract_text_from_file(filename, data)
    return structure_cv(text)
